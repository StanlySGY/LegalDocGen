from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from typing import Optional
from sqlalchemy.orm import Session

from backend.models.case import Case
from backend.models.workflow import StageType, STAGE_NAMES
from backend.services.workflow_engine.engine import WorkflowEngine


class ExportService:
    def __init__(self, db: Session):
        self.db = db
        self.engine = WorkflowEngine(db)

    def export_to_word(self, case_id: str, filename: Optional[str] = None) -> bytes:
        case = self.db.query(Case).filter(Case.id == case_id).first()
        if not case:
            raise ValueError(f"Case {case_id} not found")

        doc = Document()
        title = doc.add_heading(case.name, level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        self._add_review_notice(doc)
        self._add_case_info(doc, case)
        self._add_material_catalog(doc, case_id)
        self._add_fact_timeline(doc, case_id)

        if case.description:
            doc.add_heading("案件描述", level=2)
            self._add_text_block(doc, case.description)

        doc.add_heading("工作流输出", level=2)
        for stage_type in [
            StageType.FACT_EXTRACTION,
            StageType.LEGAL_ANALYSIS,
            StageType.DISPUTE_FOCUS,
            StageType.DRAFT_GENERATION,
            StageType.REVIEW_OPTIMIZATION,
        ]:
            node = self.engine.get_stage_node(case_id, stage_type)
            if not node or not node.output:
                continue
            doc.add_heading(STAGE_NAMES.get(stage_type, stage_type.value), level=3)
            self._add_text_block(doc, node.output)
            self._add_metadata(doc, node)

        doc.add_paragraph()
        footer = doc.add_paragraph("本文档由 LegalDocGen 辅助生成，请由专业律师复核后使用。")
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in footer.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

        from io import BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_review_notice(self, doc: Document):
        doc.add_heading("使用提示", level=2)
        paragraph = doc.add_paragraph()
        paragraph.add_run("本文件由 AI 根据上传材料和工作流结果辅助生成，不构成最终法律意见；事实、证据、法律依据和诉讼策略均需人工复核。")

    def _add_case_info(self, doc: Document, case: Case):
        table = doc.add_table(rows=4, cols=2)
        table.style = "Light Grid Accent 1"
        info_data = [
            ("案件ID", case.id),
            ("案件类型", case.case_type or "未分类"),
            ("案件状态", self._get_status_text(case.status)),
            ("导出时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ]
        for index, (label, value) in enumerate(info_data):
            table.rows[index].cells[0].text = label
            table.rows[index].cells[1].text = str(value)

    def _add_material_catalog(self, doc: Document, case_id: str):
        catalog = self.engine.get_material_catalog(case_id)
        if not catalog:
            return
        doc.add_heading("证据材料目录", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["序号", "材料名称", "类型", "解析状态", "引用页码", "内容摘要"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for index, item in enumerate(catalog, start=1):
            cells = table.add_row().cells
            cells[0].text = str(index)
            cells[1].text = item["filename"]
            cells[2].text = item["file_type"]
            cells[3].text = "已解析" if item["parse_status"] == "completed" else "失败"
            cells[4].text = item.get("citation", "页码未识别")
            cells[5].text = item["excerpt"][:120]

    def _add_fact_timeline(self, doc: Document, case_id: str):
        timeline = self.engine.get_fact_timeline(case_id)
        if not timeline:
            return
        doc.add_heading("材料事实时间线", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        for index, header in enumerate(["时间", "事件摘录", "来源材料"]):
            table.rows[0].cells[index].text = header
        for item in timeline[:30]:
            cells = table.add_row().cells
            cells[0].text = item["date"]
            cells[1].text = item["event"]
            cells[2].text = item["source"]

    def _add_text_block(self, doc: Document, text: str):
        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=4)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=4)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=3)
            else:
                paragraph = doc.add_paragraph(line)
                paragraph.paragraph_format.space_after = Pt(4)

    def _add_metadata(self, doc: Document, node):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"模型: {node.model_used or '未记录'} | ").font.size = Pt(9)
        paragraph.add_run(f"版本: {node.version} | ").font.size = Pt(9)
        paragraph.add_run(f"生成时间: {node.created_at.strftime('%Y-%m-%d %H:%M:%S') if node.created_at else '未记录'}").font.size = Pt(9)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _get_status_text(self, status: str) -> str:
        status_map = {
            "draft": "草稿",
            "in_progress": "进行中",
            "completed": "已完成",
        }
        return status_map.get(str(status), str(status))
