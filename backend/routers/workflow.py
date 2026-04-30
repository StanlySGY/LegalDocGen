import json
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse, Response
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional

from backend.database import get_db
from backend.models.workflow import StageType, STAGE_NAMES, STAGE_ORDER
from backend.models.material import Material
from backend.services.workflow_engine.engine import WorkflowEngine
from backend.services.workflow_engine.stages import STAGE_PROMPTS
from backend.services.model_dispatcher.dispatcher import dispatcher
from backend.services.prompt_manager.manager import PromptManager

router = APIRouter(prefix="/api/workflow", tags=["workflow"])


class GenerateRequest(BaseModel):
    stage: str
    prompt: str = ""
    provider: str = ""
    model: str = ""
    template_id: str = ""
    document_type: str = ""


class RollbackRequest(BaseModel):
    node_id: str


@router.get("/progress/{case_id}")
def get_progress(case_id: str, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    return engine.get_stage_progress(case_id)


@router.get("/node/{case_id}/{stage}")
def get_node(case_id: str, stage: str, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    node = engine.get_stage_node(case_id, StageType(stage))
    if not node:
        return {"stage": stage, "output": "", "prompt": STAGE_PROMPTS.get(StageType(stage), ""), "version": 0}
    return {
        "id": node.id,
        "stage": node.stage,
        "output": node.output,
        "prompt": node.prompt,
        "model_used": node.model_used,
        "version": node.version,
        "status": node.status,
        "created_at": node.created_at.isoformat() if node.created_at else None,
    }


@router.post("/generate/{case_id}")
async def generate(case_id: str, req: GenerateRequest, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    pm = PromptManager(db)
    stage = StageType(req.stage)

    progress = engine.get_stage_progress(case_id)
    stage_info = next((s for s in progress if s["stage"] == req.stage), None)
    if stage_info and stage_info["locked"]:
        raise HTTPException(400, stage_info["locked_reason"])

    prompt_template = req.prompt or pm.get_prompt(stage, req.template_id, req.document_type)
    materials_context = engine.get_case_context(case_id)
    previous_context = engine.get_previous_stages_output(case_id, stage)

    final_prompt = prompt_template.format(
        materials=materials_context["materials"],
        previous_context=previous_context,
    )
    try:
        output = await dispatcher.generate(final_prompt, req.provider, req.model)
    except Exception as e:
        raise HTTPException(500, f"生成失败: {e}")

    node = engine.create_or_update_node(
        case_id=case_id, stage=stage, prompt=req.prompt or prompt_template,
        output=output, model_used=f"{req.provider}/{req.model}" if req.provider else "default",
    )

    if stage == StageType.FACT_EXTRACTION:
        from backend.services.structurer.structurer import structure_facts
        structured = structure_facts(output)
        for m in db.query(Material).filter(Material.case_id == case_id).all():
            m.structured_data = json.dumps(structured, ensure_ascii=False)
        db.commit()

    return {"node_id": node.id, "output": output, "version": node.version}


@router.post("/generate-stream/{case_id}")
async def generate_stream(case_id: str, req: GenerateRequest, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    pm = PromptManager(db)
    stage = StageType(req.stage)

    progress = engine.get_stage_progress(case_id)
    stage_info = next((s for s in progress if s["stage"] == req.stage), None)
    if stage_info and stage_info["locked"]:
        raise HTTPException(400, stage_info["locked_reason"])

    prompt_template = req.prompt or pm.get_prompt(stage, req.template_id, req.document_type)
    materials_context = engine.get_case_context(case_id)
    previous_context = engine.get_previous_stages_output(case_id, stage)

    final_prompt = prompt_template.format(
        materials=materials_context["materials"],
        previous_context=previous_context,
    )
    async def stream_generator():
        try:
            full_output = []
            async for chunk in dispatcher.generate_stream(final_prompt, req.provider, req.model):
                full_output.append(chunk)
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"
            output = "".join(full_output)
            engine.create_or_update_node(
                case_id=case_id, stage=stage, prompt=req.prompt or prompt_template,
                output=output, model_used=f"{req.provider}/{req.model}" if req.provider else "default",
            )
            if stage == StageType.FACT_EXTRACTION:
                from backend.services.structurer.structurer import structure_facts
                structured = structure_facts(output)
                for m in db.query(Material).filter(Material.case_id == case_id).all():
                    m.structured_data = json.dumps(structured, ensure_ascii=False)
                db.commit()
            yield f"data: {json.dumps({'done': True})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(stream_generator(), media_type="text/event-stream")


@router.post("/rollback/{case_id}")
def rollback(case_id: str, req: RollbackRequest, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    node = engine.rollback_to_version(req.node_id)
    return {"stage": node.stage, "version": node.version, "output": node.output}


@router.get("/history/{case_id}/{stage}")
def get_history(case_id: str, stage: str, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    nodes = engine.get_version_history(case_id, StageType(stage))
    return [
        {
            "id": n.id, "version": n.version, "output": n.output,
            "prompt": n.prompt, "model_used": n.model_used,
            "created_at": n.created_at.isoformat() if n.created_at else None,
        }
        for n in nodes
    ]


@router.post("/save-output/{case_id}/{stage}")
def save_output(case_id: str, stage: str, output: str = "", db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    node = engine.get_stage_node(case_id, StageType(stage))
    if not node:
        raise HTTPException(404, "节点不存在")
    node.output = output
    db.commit()
    return {"message": "已保存"}


class ReviewChainRequest(BaseModel):
    models: list[dict]
    prompt: str = ""


class MultiCompareRequest(BaseModel):
    models: list[dict]
    prompt: str = ""


@router.post("/review-chain/{case_id}")
async def review_chain(case_id: str, req: ReviewChainRequest, db: Session = Depends(get_db)):
    if len(req.models) != 3:
        raise HTTPException(400, "链式审查需要3个模型（生成/审查/优化）")

    engine = WorkflowEngine(db)
    pm = PromptManager(db)
    stage = StageType.REVIEW_OPTIMIZATION
    prompt_template = req.prompt or pm.get_prompt(stage)
    materials_context = engine.get_case_context(case_id)
    previous_context = engine.get_previous_stages_output(case_id, stage)
    context_prompt = prompt_template.format(
        materials=materials_context["materials"],
        previous_context=previous_context,
    )

    from backend.services.review_orchestrator.orchestrator import ReviewOrchestrator
    orchestrator = ReviewOrchestrator()

    async def stream():
        step_outputs = {}
        final_output = ""
        try:
            async for event in orchestrator.review_chain(case_id, req.models, context_prompt):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                if event.get("step") and event.get("status") == "done":
                    step_outputs[event["step"]] = event["output"]
                if event.get("final"):
                    final_output = event.get("output", "")
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            return

        engine.create_or_update_node(
            case_id=case_id, stage=stage, prompt=req.prompt or prompt_template,
            output=final_output, model_used="review_chain",
        )
        from backend.models.review import ReviewResult
        db.add(ReviewResult(
            case_id=case_id, review_mode="chain",
            step_outputs=json.dumps(step_outputs, ensure_ascii=False),
            final_output=final_output, status="completed",
        ))
        db.commit()
        yield f"data: {json.dumps({'all_done': True})}\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")


@router.post("/multi-compare/{case_id}")
async def multi_compare(case_id: str, req: MultiCompareRequest, db: Session = Depends(get_db)):
    if len(req.models) < 2:
        raise HTTPException(400, "多版本对比需要至少2个模型")

    engine = WorkflowEngine(db)
    pm = PromptManager(db)
    stage = StageType.REVIEW_OPTIMIZATION
    prompt_template = req.prompt or pm.get_prompt(stage)
    materials_context = engine.get_case_context(case_id)
    previous_context = engine.get_previous_stages_output(case_id, stage)
    context_prompt = prompt_template.format(
        materials=materials_context["materials"],
        previous_context=previous_context,
    )

    from backend.services.review_orchestrator.orchestrator import ReviewOrchestrator
    orchestrator = ReviewOrchestrator()

    async def stream():
        try:
            async for event in orchestrator.multi_compare(case_id, req.models, context_prompt):
                yield f"data: {json.dumps(event, ensure_ascii=False)}\n\n"
                if event.get("final"):
                    model_outputs = event.get("outputs", {})
                    from backend.models.review import ReviewResult
                    db.add(ReviewResult(
                        case_id=case_id, review_mode="compare",
                        model_outputs=json.dumps(model_outputs, ensure_ascii=False),
                        status="completed",
                    ))
                    db.commit()
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")


@router.post("/review-select/{case_id}")
def review_select(case_id: str, data: dict, db: Session = Depends(get_db)):
    review_id = data.get("review_id")
    selected_model = data.get("selected_model", "")
    if not review_id:
        raise HTTPException(400, "review_id 必填")

    from backend.models.review import ReviewResult
    rr = db.query(ReviewResult).filter(ReviewResult.id == review_id).first()
    if not rr:
        raise HTTPException(404, "审查记录不存在")

    outputs = json.loads(rr.model_outputs) if rr.model_outputs else {}
    selected_output = outputs.get(selected_model, "")
    if not selected_output:
        raise HTTPException(400, f"未找到模型 {selected_model} 的输出")

    rr.selected_model = selected_model
    rr.final_output = selected_output

    engine = WorkflowEngine(db)
    engine.create_or_update_node(
        case_id=case_id, stage=StageType.REVIEW_OPTIMIZATION,
        prompt="multi_compare", output=selected_output,
        model_used=f"compare:{selected_model}",
    )
    db.commit()
    return {"output": selected_output, "model": selected_model}


class AIEditRequest(BaseModel):
    text: str
    instruction: str = ""
    provider: str = ""
    model: str = ""


@router.post("/ai-edit")
async def ai_edit(req: AIEditRequest):
    instruction = req.instruction or "润色以下法律文书文本，使其更加专业、严谨、符合法律文书的行文规范，保持原意不变："
    prompt = f"{instruction}\n\n---\n\n{req.text}"
    try:
        result = await dispatcher.generate(prompt, req.provider, req.model)
        return {"result": result}
    except Exception as e:
        raise HTTPException(500, f"AI编辑失败: {e}")


class QuickGenerateRequest(BaseModel):
    document_type: str = ""
    provider: str = ""
    model: str = ""


@router.post("/quick-generate/{case_id}")
async def quick_generate(case_id: str, req: QuickGenerateRequest, db: Session = Depends(get_db)):
    engine = WorkflowEngine(db)
    pm = PromptManager(db)
    total = len(STAGE_ORDER)

    async def run():
        for i, stage in enumerate(STAGE_ORDER):
            yield f"data: {json.dumps({'stage': stage.value, 'name': STAGE_NAMES[stage], 'status': 'running', 'progress': int(i / total * 100)}, ensure_ascii=False)}\n\n"
            doc_type = req.document_type if stage in (StageType.DRAFT_GENERATION, StageType.REVIEW_OPTIMIZATION) else ""
            prompt_template = pm.get_prompt(stage, document_type=doc_type)
            materials_context = engine.get_case_context(case_id)
            previous_context = engine.get_previous_stages_output(case_id, stage)
            try:
                final_prompt = prompt_template.format(
                    materials=materials_context["materials"],
                    previous_context=previous_context,
                )
                output = await dispatcher.generate(final_prompt, req.provider, req.model)
            except Exception as e:
                yield f"data: {json.dumps({'error': f'{STAGE_NAMES[stage]}失败: {e}'}, ensure_ascii=False)}\n\n"
                return
            node = engine.create_or_update_node(
                case_id=case_id, stage=stage, prompt=prompt_template,
                output=output, model_used=f"{req.provider}/{req.model}" if req.provider else "auto",
            )
            if stage == StageType.FACT_EXTRACTION:
                from backend.services.structurer.structurer import structure_facts
                structured = structure_facts(output)
                for m in db.query(Material).filter(Material.case_id == case_id).all():
                    m.structured_data = json.dumps(structured, ensure_ascii=False)
                db.commit()
            yield f"data: {json.dumps({'stage': stage.value, 'name': STAGE_NAMES[stage], 'status': 'done', 'progress': int((i + 1) / total * 100)}, ensure_ascii=False)}\n\n"
        final_node = engine.get_stage_node(case_id, StageType.REVIEW_OPTIMIZATION)
        final_output = final_node.output if final_node else ""
        yield f"data: {json.dumps({'done': True, 'output': final_output, 'document_type': req.document_type}, ensure_ascii=False)}\n\n"

    return StreamingResponse(run(), media_type="text/event-stream")


class ExportRequest(BaseModel):
    content: str = ""
    include_cover: bool = False
    font_size: int = 16
    margin: str = "standard"  # standard | narrow | wide
    content: str = ""


@router.post("/export/{case_id}")
def export_docx(case_id: str, req: ExportRequest, db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re

    content = req.content
    if not content:
        node = WorkflowEngine(db).get_stage_node(case_id, StageType.DRAFT_GENERATION)
        if not node or not node.output:
            raise HTTPException(404, "无文书内容可导出")
        content = node.output

    doc = Document()

    # Page margins based on margin option
    margins = {"narrow": (2.0, 2.0, 2.0, 2.0), "wide": (4.0, 4.0, 3.5, 3.5)}
    mt, mb, ml, mr = margins.get(req.margin, (3.7, 3.5, 2.8, 2.6))
    for section in doc.sections:
        section.top_margin = Cm(mt)
        section.bottom_margin = Cm(mb)
        section.left_margin = Cm(ml)
        section.right_margin = Cm(mr)

    def set_font(run, name_ascii: str, name_eastasia: str, size: Pt, bold: bool = False):
        run.font.name = name_ascii
        run.font.size = size
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name_eastasia)

    def set_paragraph(paragraph, alignment=None, first_line_indent=None, line_spacing=None, space_before=None, space_after=None):
        if alignment is not None:
            paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        if first_line_indent is not None:
            pf.first_line_indent = first_line_indent
        if line_spacing is not None:
            pf.line_spacing = line_spacing
        if space_before is not None:
            pf.space_before = space_before
        if space_after is not None:
            pf.space_after = space_after

    # Normal style defaults
    fs = Pt(req.font_size)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = fs
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
    style.paragraph_format.line_spacing = Pt(28)

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue

        # H1: document title - 黑体, centered
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_font(run, 'Times New Roman', '黑体', Pt(req.font_size + 6), bold=True)
            set_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=Pt(28), space_before=Pt(10), space_after=Pt(10))

        # H2: section title - 黑体, bold
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            set_font(run, 'Times New Roman', '黑体', fs, bold=True)
            set_paragraph(p, line_spacing=Pt(28), space_before=Pt(6), space_after=Pt(3))

        # H3: subsection - 楷体, bold
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            set_font(run, 'Times New Roman', '楷体_GB2312', fs, bold=True)
            set_paragraph(p, line_spacing=Pt(28))

        # List items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_font(run, 'Times New Roman', '仿宋_GB2312', fs)
            set_paragraph(p, first_line_indent=Cm(0.74), line_spacing=Pt(28))

        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^\d+\.\s', '', line))
            set_font(run, 'Times New Roman', '仿宋_GB2312', fs)
            set_paragraph(p, first_line_indent=Cm(0.74), line_spacing=Pt(28))

        # Bold line
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            set_font(run, 'Times New Roman', '黑体', fs, bold=True)
            set_paragraph(p, first_line_indent=Cm(0.74), line_spacing=Pt(28))

        # Normal text
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run, 'Times New Roman', '仿宋_GB2312', fs)
            set_paragraph(p, first_line_indent=Cm(0.74), line_spacing=Pt(28))

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=document.docx"},
    )
